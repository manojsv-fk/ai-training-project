# filepath: market-research-platform/backend/core/reports/export.py
# Report export service. Converts markdown report content to PDF or Word (.docx).

import io
import re
import logging

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
)
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def to_pdf(report_title: str, report_content: str) -> bytes:
    """
    Convert markdown report content to a PDF byte stream.

    Args:
        report_title: Used as the PDF title/header.
        report_content: Markdown string to render.

    Returns:
        bytes: PDF content as bytes (for StreamingResponse).
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    # Custom styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ReportTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=20,
        textColor=HexColor("#1e293b"),
    ))
    styles.add(ParagraphStyle(
        name="ReportH2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=16,
        textColor=HexColor("#334155"),
    ))
    styles.add(ParagraphStyle(
        name="ReportH3",
        parent=styles["Heading3"],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=12,
        textColor=HexColor("#475569"),
    ))
    styles.add(ParagraphStyle(
        name="ReportBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
        leading=14,
    ))
    styles.add(ParagraphStyle(
        name="ReportBullet",
        parent=styles["Normal"],
        fontSize=10,
        leftIndent=20,
        spaceAfter=4,
        leading=14,
        bulletIndent=10,
    ))

    elements = []

    # Title
    elements.append(Paragraph(report_title, styles["ReportTitle"]))
    elements.append(HRFlowable(width="100%", color=HexColor("#e2e8f0")))
    elements.append(Spacer(1, 12))

    # Parse markdown into reportlab elements
    lines = report_content.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            elements.append(Spacer(1, 6))
            continue

        # H2 headers
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            elements.append(Paragraph(text, styles["ReportH2"]))

        # H3 headers
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            elements.append(Paragraph(text, styles["ReportH3"]))

        # H1 headers (after title, treat as H2)
        elif stripped.startswith("# "):
            text = stripped[2:].strip()
            elements.append(Paragraph(text, styles["ReportH2"]))

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = _convert_markdown_inline(text)
            elements.append(Paragraph(f"• {text}", styles["ReportBullet"]))

        # Numbered list items
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped).strip()
            text = _convert_markdown_inline(text)
            elements.append(Paragraph(text, styles["ReportBullet"]))

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            elements.append(Spacer(1, 6))
            elements.append(HRFlowable(width="100%", color=HexColor("#e2e8f0")))
            elements.append(Spacer(1, 6))

        # Regular paragraph
        else:
            text = _convert_markdown_inline(stripped)
            elements.append(Paragraph(text, styles["ReportBody"]))

    # Build the PDF
    doc.build(elements)
    return buffer.getvalue()


def to_docx(report_title: str, report_content: str) -> bytes:
    """
    Convert markdown report content to a Word .docx byte stream.

    Args:
        report_title: Used as the document heading.
        report_content: Markdown string to convert.

    Returns:
        bytes: .docx content as bytes (for StreamingResponse).
    """
    doc = DocxDocument()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(report_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a thin line after the title
    doc.add_paragraph("_" * 60)

    # Parse markdown content
    lines = report_content.split("\n")
    current_paragraph = None

    for line in lines:
        stripped = line.strip()

        if not stripped:
            current_paragraph = None
            continue

        # H2 headers
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            current_paragraph = None

        # H3 headers
        elif stripped.startswith("### "):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            current_paragraph = None

        # H1 headers
        elif stripped.startswith("# "):
            text = stripped[2:].strip()
            doc.add_heading(text, level=1)
            current_paragraph = None

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = _strip_markdown_inline(text)
            para = doc.add_paragraph(text, style="List Bullet")
            _apply_inline_formatting(para, stripped[2:].strip())
            current_paragraph = None

        # Numbered list items
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped).strip()
            text = _strip_markdown_inline(text)
            para = doc.add_paragraph(text, style="List Number")
            current_paragraph = None

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            current_paragraph = None

        # Regular paragraph
        else:
            text = _strip_markdown_inline(stripped)
            para = doc.add_paragraph()
            _apply_inline_formatting(para, stripped)
            current_paragraph = para

    # Add footer with attribution
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by Market Research Intelligence Platform")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _convert_markdown_inline(text: str) -> str:
    """Convert basic markdown inline formatting to ReportLab XML tags."""
    # Bold: **text** → <b>text</b>
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic: *text* → <i>text</i>
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # Inline code: `text` → <font face="Courier">text</font>
    text = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', text)
    return text


def _strip_markdown_inline(text: str) -> str:
    """Remove markdown inline formatting for plain text output."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _apply_inline_formatting(paragraph, markdown_text: str):
    """Apply bold/italic formatting to a python-docx paragraph from markdown text."""
    # Clear existing runs and rebuild with formatting
    # This is simplified — handles **bold** and *italic* patterns
    paragraph.clear()

    # Split on bold patterns first
    parts = re.split(r"(\*\*.*?\*\*)", markdown_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic within non-bold parts
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    if sub:
                        paragraph.add_run(sub)
